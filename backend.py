from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from hugchat import hugchat
from hugchat.login import Login
import whisper
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

app = FastAPI()

EMAIL = ""
PASSWD = ""
cookie_path_dir = "./cookies/"

try:
    sign = Login(EMAIL, PASSWD)
    cookies = sign.login(cookie_dir_path=cookie_path_dir, save_cookies=True)
    chatbot = hugchat.ChatBot(cookies=cookies.get_dict())
except Exception as e:
    raise Exception(f"Failed to initialize chatbot: {e}")

class ChatRequest(BaseModel):
    path: str

# Add CORS middleware
origins = [
    "http://localhost:3000",
    "http://localhost:8000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
def format_meeting_minutes(text):
    document = Document()
    document.add_heading('Minutes of Meeting', level=1)
    sections = re.split(r'# (.+):', text)  

    for i in range(len(sections)):
        if i % 2 == 1:
            heading = sections[i]
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading + ":")
            run.bold = True
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            paragraph = document.add_paragraph(sections[i].strip())
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    return document

temp_dir = "./temp"
os.makedirs(temp_dir, exist_ok=True)

@app.post("/process/")
async def process_audio(file: UploadFile = File(...)):
    try:
        file_path = os.path.join(temp_dir, file.filename)
        print("Saving file to:", file_path)
        
        with open(file_path, "wb") as f:
            f.write(await file.read())

        print("File saved:", file_path)
        
        model = whisper.load_model("base")
        result = model.transcribe(file_path, language="en")

        print("Transcription result:", result["text"])

        
        prompt_mom = f"Generate minutes of meeting for {result['text']} with subheadings"
        mom_result = str(chatbot.query(prompt_mom, web_search=True))
        print("Minutes of meeting:", mom_result)

        prompt_summary = f"Generate summary for {result['text']}"
        summary_result = str(chatbot.query(prompt_summary, web_search=True))
        print("Summary:", summary_result)

        doc_mom = format_meeting_minutes(mom_result)
        mom_path = os.path.join(temp_dir, f"{os.path.splitext(file.filename)[0]}_Minutes_of_Meeting.docx")
        doc_mom.save(mom_path)

        doc_summary = Document()
        doc_summary.add_heading('Summary', level=1)
        doc_summary.add_paragraph(summary_result)
        summary_path = os.path.join(temp_dir, f"{os.path.splitext(file.filename)[0]}_Summary.docx")
        doc_summary.save(summary_path)

        os.remove(file_path)
        
        return {"minutes_of_meeting_path": mom_path, "summary_path": summary_path}
    except Exception as e:
        print("Error processing audio:", e)
        raise HTTPException(status_code=500, detail=f"Error processing audio: {e}")

@app.get("/download")
async def download_file(path: str):
    if os.path.exists(path):
        return FileResponse(path, filename=os.path.basename(path))
    else:
        raise HTTPException(status_code=404, detail="File not found")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
